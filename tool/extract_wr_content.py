#!/usr/bin/env python3
"""
extract_wr_content.py
=====================
Parse 2 source docx files and emit:
  assets/seed/wr_situations.json
  assets/seed/wr_stories.json

Usage:
  python3 tool/extract_wr_content.py
  python3 tool/extract_wr_content.py \
      --library /path/to/Career\ Situation\ Library.docx \
      --dataspec /path/to/WorkReflection_DataSpec_v3.docx \
      --out-dir assets/seed
"""

import argparse
import json
import re
import sys
from pathlib import Path

import docx  # python-docx

# ─────────────────────────── constants ───────────────────────────

LIBRARY_DEFAULT = (
    "/home/duythong/Desktop/FileTam/workreflection/Career Situation Library.docx"
)
DATASPEC_DEFAULT = (
    "/home/duythong/Desktop/FileTam/workreflection/WorkReflection_DataSpec_v3.docx"
)
OUT_DIR_DEFAULT = Path(__file__).parent.parent / "assets" / "seed"

ALL_DIMS = ["S1", "S2", "S3", "C1", "C2", "C3", "A1", "A2", "A3", "A4"]

# wave per plan (Fable decision, do not change)
WAVE_MAP = {
    "C2": 1, "A1": 1, "A3": 1, "C1": 1,
    "A4": 2, "A2": 2, "S1": 2,
    "C3": 3, "S2": 3, "S3": 3,
}

# Mapping from DataSpec Vietnamese label → db enum value (4 values per plan)
NEED_LABEL_MAP = {
    "Rõ ràng": "ro_rang",
    "Kết nối": "ket_noi",
    "Thích nghi": "thich_nghi",
    "Phát triển": "phat_trien",
    # Old 3-value labels in Library tables (keep for backward compat)
    "Clarity": "ro_rang",
    "Connection": "ket_noi",
    "Adaptability": "thich_nghi",  # best-effort mapping; flagged in report
}

# Career stage normalization
CAREER_STAGE_MAP = {
    "0 đến 2 năm": "0_2_years",
    "3 đến 5 năm": "3_5_years",
    "6 đến 10 năm": "6_10_years",
    "trên 10 năm": "10_plus_years",
    "Leadership": "leadership",
    "Individual Contributor": "individual_contributor",
    "Mid-level": "mid_level",
    "Senior": "senior",
    "Junior": "junior",
    "Entry": "entry",
}

# Indices of mapping tables in DataSpec docx (0-based)
# confirmed by inspection: tables 4..13 → S1..A4
DATASPEC_MAPPING_TABLE_START = 4  # table index 4 = S1
# Tables 14..16 = wave tables, 17..20 = enum tables


# ─────────────────────────── helpers ─────────────────────────────

def normalize_text(s: str) -> str:
    return " ".join(s.split())


def to_array(s: str) -> list:
    if not s or not s.strip():
        return []
    return [x.strip() for x in re.split(r"[,،،]", s) if x.strip()]


# ─────────────────────────── DataSpec parser ─────────────────────

def parse_dataspec(path: str) -> dict:
    """
    Returns:
      {
        sit_text -> {
          expected_outcome: str,
          human_need: str,       # db enum
          sca_perspective: str,
          sca_dimension: str,
        }
      }
    """
    doc = docx.Document(path)
    result = {}

    # Tables 4..13 are the 10 SCA mapping tables (one per dimension)
    for i, dim in enumerate(ALL_DIMS):
        table_idx = DATASPEC_MAPPING_TABLE_START + i
        if table_idx >= len(doc.tables):
            print(f"  [WARN] DataSpec: missing table for {dim} (idx={table_idx})", file=sys.stderr)
            continue
        table = doc.tables[table_idx]
        for row in table.rows[1:]:  # skip header
            cells = [normalize_text(c.text) for c in row.cells]
            if len(cells) < 4 or not cells[0]:
                continue
            sit_text = cells[0]
            expected_outcome = cells[1]
            need_label = cells[2]
            sca_perspective = cells[3]

            human_need = NEED_LABEL_MAP.get(need_label)
            if human_need is None:
                print(
                    f"  [WARN] Unknown human_need label {need_label!r} for {dim}/{sit_text[:30]!r}",
                    file=sys.stderr,
                )

            result[sit_text] = {
                "expected_outcome": expected_outcome or None,
                "human_need": human_need,
                "sca_perspective": sca_perspective or None,
                "sca_dimension": dim,
            }

    return result


# ─────────────────────────── Library parser ──────────────────────

def parse_library(path: str, dataspec_map: dict):
    """
    Returns (situations, stories, unmapped_situations, null_field_report).

    Situations are derived from:
      1. The 10 tables in Library (story metadata: situation text, emotion/behavior tags)
      2. Enriched with DataSpec mapping

    Stories are parsed from paragraph blocks:
      - Start: paragraph matching ^([SCA]\\d-\\d{2})\\s*[|]\\s*(.+)
      - Sections: Heading 3 'Situation' | 'Story' | 'Reflection' | 'Self Reflection' | 'Aha' | 'Practice'
    """
    doc = docx.Document(path)
    paragraphs = doc.paragraphs

    # ── 1. Parse story metadata from 10 tables ──────────────────────────────
    # Table cols: Story ID | Situation | Human Need | SCA | Emotion Tags | Behavior Tags
    # 10 tables × 10 stories = 100 rows of metadata
    meta = {}  # story_id -> dict
    for table in doc.tables:
        hdr = [normalize_text(c.text) for c in table.rows[0].cells]
        if "Story ID" not in hdr:
            continue
        for row in table.rows[1:]:
            cells = [normalize_text(c.text) for c in row.cells]
            if len(cells) < 6 or not cells[0]:
                continue
            story_id = cells[0]
            situation_text = cells[1]
            need_label_raw = cells[2]
            sca = cells[3]
            emotion_tags = to_array(cells[4])
            behavior_tags = to_array(cells[5])

            human_need_raw = NEED_LABEL_MAP.get(need_label_raw)

            meta[story_id] = {
                "story_id": story_id,
                "sca_dimension": sca,
                "situation": situation_text,
                "human_need": human_need_raw,
                "emotion_tags": emotion_tags,
                "behavior_tags": behavior_tags,
            }

    # ── 2. Parse story content blocks from paragraphs ───────────────────────
    STORY_ID_RE = re.compile(r"^([SCA]\d-\d{2})\s*[|｜]\s*(.+)$")
    SECTION_HEADINGS = {
        "situation": "situation",
        "story": "story_content",
        "reflection": "reflection_question",
        "self reflection": "self_reflection",
        "self-reflection": "self_reflection",
        "aha": "aha_message",
        "practice": "practice_action",
        # Vietnamese variants
        "tình huống": "situation",
        "câu chuyện": "story_content",
        "câu hỏi phản ánh": "reflection_question",
        "tự phản ánh": "self_reflection",
        "thông điệp aha": "aha_message",
        "hành động thực hành": "practice_action",
    }

    stories_content = {}  # story_id -> {title, story_content, reflection_question, ...}
    current_id = None
    current_title = None
    current_section = None
    current_lines = []

    def flush_section():
        nonlocal current_section, current_lines
        if current_id and current_section and current_lines:
            text = "\n".join(current_lines).strip()
            if text:
                stories_content.setdefault(current_id, {}).setdefault(
                    current_section, []
                )
                stories_content[current_id][current_section].append(text)
        current_lines = []

    def flush_story():
        flush_section()
        nonlocal current_id, current_title, current_section
        current_id = None
        current_title = None
        current_section = None

    consecutive_empty = 0  # track empty paragraph runs to detect story boundaries

    for p in paragraphs:
        text = p.text.strip()

        if not text:
            consecutive_empty += 1
            # More than 2 consecutive empty lines → likely end of story block; stop section
            if consecutive_empty > 2 and current_section:
                flush_section()
                current_section = None
            continue
        consecutive_empty = 0

        style = p.style.name.lower()

        # Check if this is a story ID heading
        m = STORY_ID_RE.match(text)
        if m:
            flush_story()
            current_id = m.group(1)
            current_title = normalize_text(m.group(2))
            stories_content[current_id] = {"title": current_title}
            continue

        if current_id is None:
            continue

        # Check if this is a section heading
        if "heading 3" in style or "heading3" in style:
            flush_section()
            heading_key = text.lower().rstrip(":")
            current_section = SECTION_HEADINGS.get(heading_key)
            if current_section is None:
                # Try partial match
                for k, v in SECTION_HEADINGS.items():
                    if k in heading_key:
                        current_section = v
                        break
            continue

        # Otherwise accumulate content
        if current_section:
            current_lines.append(text)

    flush_story()

    # ── 3. Merge metadata + content ──────────────────────────────────────────
    all_ids = sorted(set(list(meta.keys()) + list(stories_content.keys())))
    stories = []
    null_field_report = {}

    for sid in all_ids:
        m_data = meta.get(sid, {})
        c_data = stories_content.get(sid, {})

        def join_lines(field):
            val = c_data.get(field)
            if not val:
                return None
            return "\n".join(val).strip() or None

        title = c_data.get("title") or sid
        situation = join_lines("situation") or m_data.get("situation")
        story_content = join_lines("story_content")
        reflection_question = join_lines("reflection_question")
        self_reflection = join_lines("self_reflection")
        aha_message = join_lines("aha_message")
        practice_action = join_lines("practice_action")

        # SCA dim from story_id prefix
        sca_dim = m_data.get("sca_dimension") or sid[:2]

        story = {
            "story_id": sid,
            "title": title,
            "sca_dimension": sca_dim,
            "human_need": m_data.get("human_need"),
            "situation": situation,
            "emotion_tags": m_data.get("emotion_tags", []),
            "behavior_tags": m_data.get("behavior_tags", []),
            "career_stages": [],
            "difficulty_level": None,
            "story_content": story_content or "",
            "reflection_question": reflection_question,
            "self_reflection": self_reflection,
            "aha_message": aha_message,
            "practice_action": practice_action,
        }
        stories.append(story)

        # Track nulls
        null_fields = [f for f in [
            "human_need", "situation", "story_content",
            "reflection_question", "self_reflection", "aha_message", "practice_action"
        ] if not story[f]]
        if null_fields:
            null_field_report[sid] = null_fields

    # Sort by story_id
    stories.sort(key=lambda s: s["story_id"])

    # ── 4. Build situations list ─────────────────────────────────────────────
    # situations from DataSpec, enriched with wave
    # Code scheme: {DIM}-sit-{02d} numbered per dim row order
    situations = []
    unmapped = []

    # group dataspec entries by dim
    dim_entries = {d: [] for d in ALL_DIMS}
    for sit_text, info in dataspec_map.items():
        dim = info["sca_dimension"]
        dim_entries[dim].append((sit_text, info))

    for dim in ALL_DIMS:
        entries = dim_entries[dim]
        for idx, (sit_text, info) in enumerate(entries, start=1):
            code = f"{dim}-sit-{idx:02d}"
            situations.append({
                "code": code,
                "text": sit_text,
                "sca_dimension": dim,
                "human_need": info["human_need"],
                "expected_outcome": info["expected_outcome"],
                "sca_perspective": info["sca_perspective"],
                "wave": WAVE_MAP.get(dim, 3),
            })

    # Also check for situations in Library that are NOT in DataSpec
    library_situations = {}  # dim -> [text]
    for table in doc.tables:
        hdr = [normalize_text(c.text) for c in table.rows[0].cells]
        if "Story ID" in hdr:
            dim_col = hdr.index("SCA") if "SCA" in hdr else 3
            sit_col = hdr.index("Situation") if "Situation" in hdr else 1
            for row in table.rows[1:]:
                cells = [normalize_text(c.text) for c in row.cells]
                if len(cells) > max(dim_col, sit_col):
                    dim = cells[dim_col]
                    sit = cells[sit_col]
                    library_situations.setdefault(dim, []).append(sit)

    dataspec_texts = set(dataspec_map.keys())
    for dim, sit_list in library_situations.items():
        for sit in sit_list:
            if sit not in dataspec_texts:
                unmapped.append({"sca_dimension": dim, "situation": sit})

    return situations, stories, unmapped, null_field_report


# ─────────────────────────── validation ──────────────────────────

def validate(stories: list, situations: list):
    errors = []

    # Exactly 100 unique story IDs
    ids = [s["story_id"] for s in stories]
    if len(ids) != len(set(ids)):
        dup = [x for x in ids if ids.count(x) > 1]
        errors.append(f"Duplicate story IDs: {sorted(set(dup))}")
    if len(stories) != 100:
        errors.append(f"Expected 100 stories, got {len(stories)}")

    # 10 per dimension
    for dim in ALL_DIMS:
        count = sum(1 for s in stories if s["sca_dimension"] == dim)
        if count != 10:
            errors.append(f"Dimension {dim}: expected 10 stories, got {count}")

    # story_content not empty
    empty_content = [s["story_id"] for s in stories if not s.get("story_content")]
    if empty_content:
        errors.append(f"Stories with empty story_content: {empty_content}")

    return errors


# ─────────────────────────── main ────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Extract WR content to JSON seeds")
    parser.add_argument("--library", default=LIBRARY_DEFAULT)
    parser.add_argument("--dataspec", default=DATASPEC_DEFAULT)
    parser.add_argument("--out-dir", default=str(OUT_DIR_DEFAULT))
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 60)
    print("WR Content Extraction")
    print("=" * 60)
    print(f"Library  : {args.library}")
    print(f"DataSpec : {args.dataspec}")
    print(f"Out dir  : {out_dir}")
    print()

    print("Step 1: Parsing DataSpec v3...")
    dataspec_map = parse_dataspec(args.dataspec)
    print(f"  → {len(dataspec_map)} situation mappings extracted from DataSpec")

    print("\nStep 2: Parsing Career Situation Library...")
    situations, stories, unmapped, null_report = parse_library(
        args.library, dataspec_map
    )

    print(f"  → {len(stories)} stories extracted")
    print(f"  → {len(situations)} situations built from DataSpec")

    print("\nStep 3: Validation...")
    errors = validate(stories, situations)
    if errors:
        print("  [ERROR] Validation failed:")
        for e in errors:
            print(f"    - {e}")
    else:
        print("  [OK] All validation checks passed")

    print("\nStep 4: Writing JSON output...")
    sit_path = out_dir / "wr_situations.json"
    sto_path = out_dir / "wr_stories.json"

    with open(sit_path, "w", encoding="utf-8") as f:
        json.dump(situations, f, ensure_ascii=False, indent=2)
    with open(sto_path, "w", encoding="utf-8") as f:
        json.dump(stories, f, ensure_ascii=False, indent=2)

    print(f"  Written: {sit_path}")
    print(f"  Written: {sto_path}")

    print("\n" + "=" * 60)
    print("REPORT")
    print("=" * 60)

    print(f"\nTotal situations seeded : {len(situations)}")
    print(f"Total stories seeded    : {len(stories)}")

    print("\nDistribution by SCA dimension:")
    for dim in ALL_DIMS:
        n_stories = sum(1 for s in stories if s["sca_dimension"] == dim)
        n_sits = sum(1 for s in situations if s["sca_dimension"] == dim)
        wave = WAVE_MAP.get(dim, 3)
        print(f"  {dim} (wave{wave}): {n_stories} stories, {n_sits} situations")

    if unmapped:
        print(f"\nSituations in Library WITHOUT DataSpec mapping ({len(unmapped)}):")
        for u in unmapped:
            print(f"  [{u['sca_dimension']}] {u['situation']}")
    else:
        print("\nAll Library situations found in DataSpec mapping.")

    if null_report:
        print(f"\nNull/empty fields by story ({len(null_report)} stories affected):")
        # Group by field
        field_counts = {}
        for sid, fields in null_report.items():
            for f in fields:
                field_counts[f] = field_counts.get(f, 0) + 1
        for field, cnt in sorted(field_counts.items(), key=lambda x: -x[1]):
            print(f"  {field}: {cnt} stories missing")
        print("\n  Per-story detail:")
        for sid, fields in sorted(null_report.items()):
            print(f"    {sid}: {fields}")
    else:
        print("\nNo null/empty fields in stories.")

    if errors:
        print(f"\n[FAIL] {len(errors)} validation error(s). Exiting with code 1.")
        sys.exit(1)
    else:
        print("\n[PASS] Extraction complete.")


if __name__ == "__main__":
    main()
